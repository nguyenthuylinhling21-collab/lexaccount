import streamlit as st
import fitz
from docx import Document
import os
import requests
import pytesseract
from PIL import Image
from pdf2image import convert_from_path

OPENROUTER_API_KEY = "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"
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
POPPLER_PATH = r"C:\poppler\poppler\Library\bin"

st.set_page_config(page_title="LexAccount", page_icon="⚖️", layout="wide")

st.markdown("""
<style>
* { font-family: 'Open Sans', sans-serif; }
[data-testid="stAppViewContainer"] { background-color: #f5f5f5; }
[data-testid="stSidebar"] { background-color: #000000; }
[data-testid="stSidebar"] * { color: white !important; }
.stButton > button { background-color: #86BC25; color: white; border-radius: 0px; border: none; padding: 10px 28px; font-weight: 700; text-transform: uppercase; font-size: 13px; }
.stButton > button:hover { background-color: #6a9a1e; }
.stTabs [data-baseweb="tab-list"] { gap: 0px; border-bottom: 3px solid #86BC25; }
.stTabs [data-baseweb="tab"] { background-color: white; border-radius: 0px; padding: 10px 24px; font-weight: 600; color: #000000; border: 1px solid #ddd; }
.stTabs [aria-selected="true"] { background-color: #86BC25 !important; color: white !important; border: none !important; }
.stTextInput input { border-radius: 0px; border: 1px solid #cccccc; border-bottom: 2px solid #86BC25; padding: 10px; }
.answer-box { background-color: white; border-left: 4px solid #86BC25; padding: 24px; margin-top: 16px; box-shadow: 0 1px 4px rgba(0,0,0,0.1); line-height: 1.8; }
.stat-box { background: white; border-top: 3px solid #86BC25; padding: 20px 24px; box-shadow: 0 1px 4px rgba(0,0,0,0.08); }
.stat-num { font-size: 2.2rem; font-weight: 700; color: #000000; }
.stat-label { font-size: 0.82rem; color: #666; margin-top: 4px; text-transform: uppercase; letter-spacing: 0.5px; }
.section-title { font-size: 1.1rem; font-weight: 700; color: #000000; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 12px; border-bottom: 2px solid #86BC25; padding-bottom: 6px; }
</style>
""", unsafe_allow_html=True)

os.makedirs("documents", exist_ok=True)

if "last_question" not in st.session_state:
    st.session_state.last_question = ""

def ask_ai(prompt):
    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"},
        json={"model": "nvidia/nemotron-3-super-120b-a12b:free", "messages": [{"role": "user", "content": prompt}]}
    )
    result = response.json()
    if "choices" in result:
        return result["choices"][0]["message"]["content"]
    return str(result)

def doc_to_text(path, filename):
    if filename.endswith(".pdf"):
        doc = fitz.open(path)
        text = "\n".join([doc[i].get_text() for i in range(len(doc))])
        if len(text.strip()) < 100:
            st.info("Phát hiện PDF scan, đang dùng OCR...")
            try:
                images = convert_from_path(path, poppler_path=POPPLER_PATH)
                text = ""
                for img in images:
                    text += pytesseract.image_to_string(img, lang="vie+eng") + "\n"
            except Exception as e:
                st.error(f"Lỗi OCR: {str(e)}")
        return text
    elif filename.endswith(".docx"):
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return ""

def search_in_docs(keyword, files):
    results = []
    for fname in files:
        path = os.path.join("documents", fname)
        text = doc_to_text(path, fname)
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if keyword.lower() in line.lower() and line.strip():
                results.append({"file": fname, "dong": i + 1, "noi_dung": line.strip()})
    return results

with st.sidebar:
    st.markdown("## ⚖️ LexAccount")
    st.markdown("*Thư viện Pháp luật Kế toán & Tài chính*")
    st.divider()
    st.markdown("**📁 THÊM VĂN BẢN**")
    uploaded_files = st.file_uploader("Upload PDF hoặc DOCX", type=["pdf", "docx"], accept_multiple_files=True)
    if uploaded_files:
        for uploaded_file in uploaded_files:
            save_path = os.path.join("documents", uploaded_file.name)
            with open(save_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
        st.success(f"Đã lưu {len(uploaded_files)} văn bản!")
    st.divider()
    st.markdown("**📚 VĂN BẢN TRONG THƯ VIỆN**")
    files = os.listdir("documents")
    if files:
        for f in files:
            col1, col2 = st.columns([3, 1])
            col1.markdown(f"📄 {f[:20]}...")
            if col2.button("Xóa", key=f):
                os.remove(os.path.join("documents", f))
                st.rerun()
    else:
        st.info("Chưa có văn bản nào")

st.markdown("""
<div style="background:#000;padding:18px 28px;margin:-1rem -1rem 0 -1rem;display:flex;align-items:center;gap:8px">
<span style="color:white;font-size:1.6rem;font-weight:700">LexAccount</span>
<span style="color:#86BC25;font-size:2rem;font-weight:900">.</span>
</div>
""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)
st.markdown('<p style="color:#555;font-size:0.9rem;border-left:3px solid #86BC25;padding-left:12px;margin-bottom:24px">Thư viện Pháp luật Kế toán & Tài chính Doanh nghiệp — Tra cứu thông minh, trích dẫn chính xác</p>', unsafe_allow_html=True)

files = os.listdir("documents")
col1, col2, col3 = st.columns(3)
with col1:
    st.markdown(f'<div class="stat-box"><div class="stat-num">{len(files)}</div><div class="stat-label">Văn bản trong thư viện</div></div>', unsafe_allow_html=True)
with col2:
    total_size = sum(os.path.getsize(os.path.join("documents", f)) for f in files) // 1024 if files else 0
    st.markdown(f'<div class="stat-box"><div class="stat-num">{total_size} KB</div><div class="stat-label">Tổng dung lượng</div></div>', unsafe_allow_html=True)
with col3:
    st.markdown('<div class="stat-box"><div class="stat-num">AI</div><div class="stat-label">Trích dẫn tự động</div></div>', unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)
tab1, tab2, tab3 = st.tabs(["🤖  Hỏi đáp AI", "🔍  Tìm kiếm nhanh", "📖  Xem văn bản"])

with tab1:
    st.markdown('<p class="section-title">Hỏi đáp pháp luật có trích dẫn</p>', unsafe_allow_html=True)
    if not files:
        st.warning("Hãy upload văn bản vào thư viện trước.")
    else:
        selected_docs = st.multiselect("Chọn văn bản để hỏi:", files, default=files)
        question = st.text_input("Nhập câu hỏi (nhấn Enter để hỏi):", placeholder="Ví dụ: Thông tư 99 quy định gì về kế toán?", key="question_input")
        auto_submit = question.strip() != "" and question != st.session_state.last_question
        st.session_state.last_question = question
        if (st.button("🔎 HỎI AI") or auto_submit) and question.strip() and selected_docs:
            with st.spinner("Đang phân tích tài liệu..."):
                all_text = ""
                for fname in selected_docs:
                    path = os.path.join("documents", fname)
                    text = doc_to_text(path, fname)
                    all_text += f"\n\n=== TÀI LIỆU: {fname} ===\n{text[:5000]}"
                prompt = f"""Bạn là chuyên gia pháp luật kế toán tài chính Việt Nam.
Dựa vào các tài liệu pháp luật dưới đây, hãy trả lời câu hỏi.
Yêu cầu: Nêu rõ trích dẫn từ tài liệu nào, điều khoản nào cụ thể.
Trả lời bằng tiếng Việt đầy đủ dấu.

TÀI LIỆU:
{all_text}

CÂU HỎI: {question}

Trả lời chi tiết, nêu rõ nguồn trích dẫn."""
                try:
                    answer = ask_ai(prompt)
                    st.markdown(f'<div class="answer-box">{answer}</div>', unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Lỗi: {str(e)}")

with tab2:
    st.markdown('<p class="section-title">Tìm kiếm nhanh trong văn bản</p>', unsafe_allow_html=True)
    if not files:
        st.warning("Hãy upload văn bản vào thư viện trước.")
    else:
        keyword = st.text_input("Nhập từ khóa tìm kiếm:", placeholder="Ví dụ: điều 5, kế toán trưởng, hóa đơn...")
        if st.button("🔍 TÌM KIẾM") and keyword:
            with st.spinner("Đang tìm kiếm..."):
                results = search_in_docs(keyword, files)
            if results:
                st.success(f"Tìm thấy **{len(results)}** kết quả cho từ khóa: *{keyword}*")
                for r in results[:20]:
                    with st.expander(f"📄 {r['file']} — Dòng {r['dong']}"):
                        st.markdown(r['noi_dung'])
            else:
                st.info(f"Không tìm thấy kết quả nào cho: *{keyword}*")

with tab3:
    st.markdown('<p class="section-title">Xem nội dung văn bản</p>', unsafe_allow_html=True)
    if not files:
        st.warning("Hãy upload văn bản vào thư viện trước.")
    else:
        selected = st.selectbox("Chọn văn bản để xem:", files)
        if selected:
            path = os.path.join("documents", selected)
            if selected.endswith(".pdf"):
                doc = fitz.open(path)
                st.info(f"Tổng số trang: **{len(doc)}**")
                page_num = st.slider("Chọn trang:", 1, len(doc), 1)
                page = doc[page_num - 1]
                text = page.get_text()
                if len(text.strip()) < 50:
                    st.info("Đang dùng OCR để đọc trang này...")
                    try:
                        images = convert_from_path(path, first_page=page_num, last_page=page_num, poppler_path=POPPLER_PATH)
                        text = pytesseract.image_to_string(images[0], lang="vie+eng")
                    except Exception as e:
                        st.error(f"Lỗi OCR: {str(e)}")
                st.text_area("Nội dung:", text, height=400)
            elif selected.endswith(".docx"):
                doc = Document(path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                st.text_area("Nội dung:", full_text, height=400)